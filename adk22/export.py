import json
from docx import Document

# Load JSON file
with open("session-c9.json", "r", encoding="utf-8") as f:
    data = json.load(f)

# Create a new Word document
doc = Document()
doc.add_heading("Google Maps Agent Conversation", level=1)

# Iterate through events
for event in data.get("events", []):
    author = event.get("author", "")
    parts = event.get("content", {}).get("parts", [])

    if not parts:
        continue

    # Combine all text parts
    text = " ".join(part.get("text", "") for part in parts).strip()
    if not text:
        continue

    # Determine label
    if author == "user":
        label = "User:"
    elif author == "maps_assistant_agent":
        label = "Agent:"
    else:
        continue  # skip unknown authors

    # Add paragraph to doc
    doc.add_paragraph(f"{label} {text}")

# Save the document
doc.save("conversation.docx")
print("Conversation saved to conversation.docx")