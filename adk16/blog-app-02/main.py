# main.py

import asyncio
import re
import os
import warnings
import logging
from datetime import datetime

# ── Suppress ADK internal warnings ───────────────────────────────────────────
warnings.filterwarnings("ignore", message=".*non-text parts in the response.*")
warnings.filterwarnings("ignore", message=".*PROGRESSIVE_SSE_STREAMING.*", category=UserWarning)
logging.getLogger("google.adk").setLevel(logging.ERROR)
# ─────────────────────────────────────────────────────────────────────────────

from google.adk.runners import Runner
from google.adk.sessions import DatabaseSessionService
from google.adk.memory import InMemoryMemoryService
from google.genai import types
from agent import app, APP_NAME
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paths ─────────────────────────────────────────────────────────────────────
OUTPUT_DIR = "output"
DB_PATH    = "blog_agent_database.db"       # SQLite file — created automatically
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Persistent session service — uses aiosqlite (async driver) ────────────────
session_service = DatabaseSessionService(db_url=f"sqlite+aiosqlite:///./{DB_PATH}")

# ── In-memory memory service ──────────────────────────────────────────────────
memory_service = InMemoryMemoryService()

# ── Runner ────────────────────────────────────────────────────────────────────
runner = Runner(
    agent=app.root_agent,          # Runner expects agent, not app
    app_name=APP_NAME,
    session_service=session_service,
    memory_service=memory_service,
)

USER_ID = "blog_user"


# ── .docx writer ──────────────────────────────────────────────────────────────
def save_blog_to_docx(topic: str, blog_content: str) -> str:
    """Saves the blog content to a properly formatted .docx Word file."""
    safe_topic = re.sub(r"[^\w\s-]", "", topic).strip()
    safe_topic = re.sub(r"\s+", "_", safe_topic)[:60]
    now = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_topic}_{now}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    for line in blog_content.strip().split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            for run in h.runs:
                run.font.size = Pt(22)

        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                run.font.size = Pt(16)

        elif line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            for run in h.runs:
                run.font.size = Pt(13)

        elif line.startswith(("- ", "* ")):
            para = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            for run in para.runs:
                run.font.size = Pt(11)

        elif re.match(r"^\d+\.\s", line):
            para = doc.add_paragraph(
                re.sub(r"^\d+\.\s", "", line).strip(), style="List Number"
            )
            for run in para.runs:
                run.font.size = Pt(11)

        else:
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.size = Pt(11)

    doc.save(filepath)
    return filepath


# ── Session helper ────────────────────────────────────────────────────────────
async def get_or_create_session(topic: str):
    """
    Reuse an existing SQLite session for this topic if it exists,
    otherwise create a fresh one.  Session ID is derived from the topic
    so the same topic always resumes the same conversation thread.
    """
    # Initialise DB tables if the service supports it
    if hasattr(session_service, "initialize"):
        await session_service.initialize()

    safe_id    = re.sub(r"\W+", "_", topic.strip().lower())[:60]
    session_id = f"session_{safe_id}"

    session = await session_service.get_session(
        app_name=APP_NAME,
        user_id=USER_ID,
        session_id=session_id,
    )

    if session:
        print(f"  [memory] Resuming existing session: {session_id}")
        return session

    session = await session_service.create_session(
        app_name=APP_NAME,
        user_id=USER_ID,
        session_id=session_id,
        state={"topic": topic},
    )
    print(f"  [memory] Created new session: {session_id}")
    return session


# ── Main ──────────────────────────────────────────────────────────────────────
async def main():
    print("=" * 55)
    print("   Research Blog Generator — Powered by Google ADK")
    print("=" * 55)

    topic = input("\nEnter the blog topic: ").strip()
    if not topic:
        print("Error: Topic cannot be empty.")
        return

    print(f"\nGenerating blog for: '{topic}' ...\n")

    session = await get_or_create_session(topic)

    content = types.Content(
        role="user",
        parts=[types.Part(text=topic)],
    )

    # ── Collect all text events keyed by agent author ─────────────────────────
    all_text: dict[str, list[str]] = {}

    async for event in runner.run_async(
        user_id=USER_ID,
        session_id=session.id,
        new_message=content,
    ):
        if not event.content:
            continue
        author = event.author or "unknown"
        for part in event.content.parts:
            text = getattr(part, "text", None)
            if text and text.strip():
                all_text.setdefault(author, []).append(text.strip())

    if not all_text:
        print("No content was generated. Please try again.")
        return

    # ── Pick the best blog text ───────────────────────────────────────────────
    # Priority: blog_generator_agent → controller_agent → longest block
    blog_content = None
    for preferred in ["blog_generator_agent", "controller_agent"]:
        if preferred in all_text:
            candidate = "\n".join(all_text[preferred]).strip()
            if len(candidate) > 200:
                blog_content = candidate
                break

    if not blog_content:
        blog_content = max(
            (" ".join(parts) for parts in all_text.values()),
            key=len,
        ).strip()

    print("\n📘 Generated Blog:\n")
    print(blog_content)

    # ── Save to .docx ─────────────────────────────────────────────────────────
    filepath = save_blog_to_docx(topic, blog_content)
    print(f"\n✅ Blog saved to: {filepath}")
    print(f"💾 Session persisted in: {DB_PATH}")

    # ── Show session summary ──────────────────────────────────────────────────
    final_session = await session_service.get_session(
        app_name=APP_NAME,
        user_id=USER_ID,
        session_id=session.id,
    )
    if final_session:
        print(f"\n--- Session Summary ---")
        print(f"Session ID : {final_session.id}")
        print(f"State      : {final_session.state}")
        print(f"Events     : {len(final_session.events)}")
        print(f"-----------------------\n")


if __name__ == "__main__":
    asyncio.run(main())